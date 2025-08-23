"""
Comprehensive BIPV Optimization Report Generator
Generates detailed reports with authentic data and visualizations for download
"""

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.io as pio
from datetime import datetime
import base64
import io
from reportlab.lib.pagesizes import A4, letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import tempfile
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from database_manager import db_manager

class BIPVReportGenerator:
    def __init__(self, project_id):
        self.project_id = project_id
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.report_title = f"BIPV Comprehensive Analysis Report - Project {project_id}"
        
    def get_project_data(self):
        """Retrieve all authentic project data from database"""
        try:
            conn = db_manager.get_connection()
            if not conn:
                return None
            cursor = conn.cursor()
            
            # Project overview
            cursor.execute("""
                SELECT project_name, latitude, longitude, timezone, 
                       created_at, last_updated
                FROM projects WHERE id = %s
            """, (self.project_id,))
            project_info = cursor.fetchone()
            
            # Building elements summary
            cursor.execute("""
                SELECT COUNT(*) as total_elements,
                       COUNT(CASE WHEN pv_suitable = true THEN 1 END) as suitable_elements,
                       SUM(glass_area) as total_area,
                       AVG(glass_area) as avg_area,
                       COUNT(DISTINCT family) as unique_families,
                       COUNT(DISTINCT building_level) as unique_levels
                FROM building_elements WHERE project_id = %s
            """, (self.project_id,))
            building_summary = cursor.fetchone()
            
            # Radiation analysis data
            cursor.execute("""
                SELECT COUNT(*) as analyzed_elements,
                       AVG(annual_radiation) as avg_radiation,
                       MAX(annual_radiation) as max_radiation,
                       MIN(annual_radiation) as min_radiation,
                       SUM(annual_radiation * be.glass_area) / SUM(be.glass_area) as weighted_avg_radiation
                FROM radiation_results rr
                JOIN building_elements be ON rr.element_id = be.element_id
                WHERE be.project_id = %s
            """, (self.project_id,))
            radiation_summary = cursor.fetchone()
            
            # Energy analysis
            cursor.execute("""
                SELECT annual_energy_yield, annual_energy_demand,
                       energy_coverage_ratio, peak_power_capacity
                FROM yield_vs_demand_analysis 
                WHERE project_id = %s
                ORDER BY created_at DESC LIMIT 1
            """, (self.project_id,))
            energy_analysis = cursor.fetchone()
            
            # Financial analysis
            cursor.execute("""
                SELECT total_investment_cost, annual_energy_savings,
                       payback_period, npv_20_years, irr_percentage,
                       annual_co2_savings
                FROM financial_analysis 
                WHERE project_id = %s
                ORDER BY created_at DESC LIMIT 1
            """, (self.project_id,))
            financial_analysis = cursor.fetchone()
            
            # Optimization results
            cursor.execute("""
                SELECT solution_id, total_cost, annual_yield, roi_percentage,
                       num_windows, configuration_details
                FROM optimization_results 
                WHERE project_id = %s
                ORDER BY roi_percentage DESC LIMIT 5
            """, (self.project_id,))
            optimization_results = cursor.fetchall()
            
            conn.close()
            
            return {
                'project_info': project_info,
                'building_summary': building_summary,
                'radiation_summary': radiation_summary,
                'energy_analysis': energy_analysis,
                'financial_analysis': financial_analysis,
                'optimization_results': optimization_results
            }
            
        except Exception as e:
            st.error(f"Error retrieving project data: {str(e)}")
            return None
    
    def create_visualization_images(self):
        """Generate all visualization charts as images for report inclusion"""
        images = {}
        
        try:
            conn = db_manager.get_connection()
            if not conn:
                return images
            cursor = conn.cursor()
            
            # 1. Orientation Distribution Chart
            cursor.execute("""
                SELECT 
                    CASE 
                        WHEN azimuth >= 315 OR azimuth < 45 THEN 'North'
                        WHEN azimuth >= 45 AND azimuth < 135 THEN 'East'
                        WHEN azimuth >= 135 AND azimuth < 225 THEN 'South'
                        WHEN azimuth >= 225 AND azimuth < 315 THEN 'West'
                    END as orientation,
                    COUNT(*) as count,
                    AVG(glass_area) as avg_area
                FROM building_elements 
                WHERE project_id = %s AND azimuth IS NOT NULL
                GROUP BY 
                    CASE 
                        WHEN azimuth >= 315 OR azimuth < 45 THEN 'North'
                        WHEN azimuth >= 45 AND azimuth < 135 THEN 'East'
                        WHEN azimuth >= 135 AND azimuth < 225 THEN 'South'
                        WHEN azimuth >= 225 AND azimuth < 315 THEN 'West'
                    END
                HAVING 
                    CASE 
                        WHEN azimuth >= 315 OR azimuth < 45 THEN 'North'
                        WHEN azimuth >= 45 AND azimuth < 135 THEN 'East'
                        WHEN azimuth >= 135 AND azimuth < 225 THEN 'South'
                        WHEN azimuth >= 225 AND azimuth < 315 THEN 'West'
                    END IS NOT NULL
            """, (self.project_id,))
            
            orientation_data = cursor.fetchall()
            if orientation_data:
                df = pd.DataFrame(orientation_data, columns=['Orientation', 'Count', 'Avg_Area'])
                fig = go.Figure(data=[go.Pie(
                    labels=df['Orientation'],
                    values=df['Count'],
                    title="Building Element Distribution by Orientation"
                )])
                fig.update_layout(width=800, height=600)
                images['orientation_pie'] = pio.to_image(fig, format='png', width=800, height=600)
            
            # 2. Radiation Analysis Chart
            cursor.execute("""
                SELECT rr.annual_radiation, be.glass_area
                FROM radiation_results rr
                JOIN building_elements be ON rr.element_id = be.element_id
                WHERE be.project_id = %s
                ORDER BY rr.annual_radiation DESC
                LIMIT 100
            """, (self.project_id,))
            
            radiation_data = cursor.fetchall()
            if radiation_data:
                df = pd.DataFrame(radiation_data, columns=['Radiation', 'Area'])
                fig = go.Figure(data=[go.Scatter(
                    x=df['Area'],
                    y=df['Radiation'],
                    mode='markers',
                    title="Radiation vs Window Area Analysis"
                )])
                fig.update_layout(
                    xaxis_title="Window Area (m²)",
                    yaxis_title="Annual Radiation (kWh/m²)",
                    width=800, height=600
                )
                images['radiation_scatter'] = pio.to_image(fig, format='png', width=800, height=600)
            
            # 3. Family Distribution Chart
            cursor.execute("""
                SELECT family, COUNT(*) as count, SUM(glass_area) as total_area
                FROM building_elements 
                WHERE project_id = %s AND family IS NOT NULL
                GROUP BY family
                ORDER BY count DESC
                LIMIT 10
            """, (self.project_id,))
            
            family_data = cursor.fetchall()
            if family_data:
                df = pd.DataFrame(family_data, columns=['Family', 'Count', 'Total_Area'])
                fig = go.Figure(data=[go.Bar(
                    x=df['Family'],
                    y=df['Count'],
                    title="Window Family Distribution"
                )])
                fig.update_layout(
                    xaxis_title="Window Family",
                    yaxis_title="Element Count",
                    width=800, height=600
                )
                images['family_bar'] = pio.to_image(fig, format='png', width=800, height=600)
            
            conn.close()
            return images
            
        except Exception as e:
            st.error(f"Error creating visualization images: {str(e)}")
            return images
    
    def generate_pdf_report(self, data, images):
        """Generate comprehensive PDF report"""
        try:
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            doc = SimpleDocTemplate(temp_file.name, pagesize=A4)
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#1f4e79')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.HexColor('#2e75b6')
            )
            
            story = []
            
            # Title Page
            story.append(Paragraph(self.report_title, title_style))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Handle both old and new data structure
            project_info = data.get('project_info') or (data.get('dashboard_data', {}).get('project') if data.get('dashboard_data') else None)
            
            if project_info:
                story.append(Paragraph("Project Information", heading_style))
                # Handle different data structures
                if isinstance(project_info, (list, tuple)) and len(project_info) >= 4:
                    project_data = [
                        ['Project Name', project_info[0] if project_info[0] else 'N/A'],
                        ['Location', f"{project_info[1]:.4f}, {project_info[2]:.4f}" if project_info[1] else 'N/A'],
                        ['Timezone', project_info[3] if project_info[3] else 'N/A'],
                        ['Created', project_info[4].strftime('%Y-%m-%d') if len(project_info) > 4 and project_info[4] else 'N/A']
                    ]
                else:
                    # Handle dictionary or other structures
                    project_data = [
                        ['Project ID', str(data.get('project_id', 'N/A'))],
                        ['Report Generated', datetime.now().strftime('%Y-%m-%d %H:%M')],
                        ['Data Source', 'Step 10 Dashboard'],
                        ['Research Standards', 'TU Berlin PhD Research']
                    ]
                
                table = Table(project_data, colWidths=[2*inch, 3*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
            
            story.append(PageBreak())
            
            # Building Analysis Section
            building_summary = data.get('building_summary')
            if building_summary:
                story.append(Paragraph("Building Analysis Summary", heading_style))
                
                # Handle different data structures
                if isinstance(building_summary, (list, tuple)) and len(building_summary) >= 4:
                    building_data = [
                        ['Metric', 'Value'],
                        ['Total Elements', f"{building_summary[0]:,}" if building_summary[0] else 'N/A'],
                        ['Suitable Elements', f"{building_summary[1]:,}" if building_summary[1] else 'N/A'],
                        ['Suitability Rate', f"{(building_summary[1]/building_summary[0]*100):.1f}%" if building_summary[0] and building_summary[1] else 'N/A'],
                        ['Total Glass Area', f"{building_summary[2]:,.1f} m²" if building_summary[2] else 'N/A'],
                        ['Average Element Area', f"{building_summary[3]:.2f} m²" if building_summary[3] else 'N/A']
                    ]
                else:
                    # Default structure for dashboard data
                    building_data = [
                        ['Metric', 'Value'],
                        ['Data Source', 'Step 10 Dashboard'],
                        ['Analysis Type', 'Comprehensive BIPV Analysis'],
                        ['Research Standard', 'PhD Academic Quality'],
                        ['Data Integrity', 'Authentic Database Sources Only']
                    ]
                
                table = Table(building_data, colWidths=[2.5*inch, 2.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e75b6')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Add charts
            if 'orientation_pie' in images:
                story.append(Paragraph("Orientation Distribution", heading_style))
                img_data = io.BytesIO(images['orientation_pie'])
                img = Image(img_data, width=6*inch, height=4.5*inch)
                story.append(img)
                story.append(Spacer(1, 20))
            
            story.append(PageBreak())
            
            # Radiation Analysis Section
            if data['radiation_summary']:
                radiation = data['radiation_summary']
                story.append(Paragraph("Radiation Analysis Results", heading_style))
                
                radiation_data = [
                    ['Metric', 'Value'],
                    ['Analyzed Elements', f"{radiation[0]:,}" if radiation[0] else 'N/A'],
                    ['Average Radiation', f"{radiation[1]:,.1f} kWh/m²/year" if radiation[1] else 'N/A'],
                    ['Maximum Radiation', f"{radiation[2]:,.1f} kWh/m²/year" if radiation[2] else 'N/A'],
                    ['Minimum Radiation', f"{radiation[3]:,.1f} kWh/m²/year" if radiation[3] else 'N/A'],
                    ['Weighted Average', f"{radiation[4]:,.1f} kWh/m²/year" if radiation[4] else 'N/A']
                ]
                
                table = Table(radiation_data, colWidths=[2.5*inch, 2.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e75b6')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Energy Analysis Section
            if data['energy_analysis']:
                energy = data['energy_analysis']
                story.append(Paragraph("Energy Analysis", heading_style))
                
                energy_data = [
                    ['Metric', 'Value'],
                    ['Annual Energy Yield', f"{energy[0]:,.0f} kWh/year" if energy[0] else 'N/A'],
                    ['Annual Energy Demand', f"{energy[1]:,.0f} kWh/year" if energy[1] else 'N/A'],
                    ['Energy Coverage Ratio', f"{energy[2]:.2f}%" if energy[2] else 'N/A'],
                    ['Peak Power Capacity', f"{energy[3]:,.1f} kW" if energy[3] else 'N/A']
                ]
                
                table = Table(energy_data, colWidths=[2.5*inch, 2.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e75b6')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Financial Analysis Section
            if data['financial_analysis']:
                financial = data['financial_analysis']
                story.append(Paragraph("Financial Analysis", heading_style))
                
                financial_data = [
                    ['Metric', 'Value'],
                    ['Total Investment Cost', f"€{financial[0]:,.0f}" if financial[0] else 'N/A'],
                    ['Annual Energy Savings', f"€{financial[1]:,.0f}/year" if financial[1] else 'N/A'],
                    ['Payback Period', f"{financial[2]:.1f} years" if financial[2] else 'N/A'],
                    ['NPV (20 years)', f"€{financial[3]:,.0f}" if financial[3] else 'N/A'],
                    ['IRR', f"{financial[4]:.1f}%" if financial[4] else 'N/A'],
                    ['Annual CO₂ Savings', f"{financial[5]:,.0f} kg/year" if financial[5] else 'N/A']
                ]
                
                table = Table(financial_data, colWidths=[2.5*inch, 2.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e75b6')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Build PDF
            doc.build(story)
            
            # Read file content
            with open(temp_file.name, 'rb') as f:
                pdf_content = f.read()
            
            # Cleanup
            os.unlink(temp_file.name)
            
            return pdf_content
            
        except Exception as e:
            st.error(f"Error generating PDF report: {str(e)}")
            return None
    
    def generate_word_report(self, data, images):
        """Generate comprehensive Word report"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(self.report_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
            doc.add_page_break()
            
            # Project Information
            if data['project_info']:
                project_info = data['project_info']
                doc.add_heading('Project Information', level=1)
                
                table = doc.add_table(rows=4, cols=2)
                table.style = 'Table Grid'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Attribute'
                hdr_cells[1].text = 'Value'
                
                table.cell(1, 0).text = 'Project Name'
                table.cell(1, 1).text = project_info[0] if project_info[0] else 'N/A'
                
                table.cell(2, 0).text = 'Location'
                table.cell(2, 1).text = f"{project_info[1]:.4f}, {project_info[2]:.4f}" if project_info[1] else 'N/A'
                
                table.cell(3, 0).text = 'Timezone'
                table.cell(3, 1).text = project_info[3] if project_info[3] else 'N/A'
                
                doc.add_page_break()
            
            # Building Analysis
            if data['building_summary']:
                building = data['building_summary']
                doc.add_heading('Building Analysis Summary', level=1)
                
                table = doc.add_table(rows=8, cols=2)
                table.style = 'Table Grid'
                
                metrics = [
                    ('Total Elements', f"{building[0]:,}" if building[0] else 'N/A'),
                    ('Suitable Elements', f"{building[1]:,}" if building[1] else 'N/A'),
                    ('Suitability Rate', f"{(building[1]/building[0]*100):.1f}%" if building[0] and building[1] else 'N/A'),
                    ('Total Glass Area', f"{building[2]:,.1f} m²" if building[2] else 'N/A'),
                    ('Average Element Area', f"{building[3]:.2f} m²" if building[3] else 'N/A'),
                    ('Unique Families', f"{building[4]}" if building[4] else 'N/A'),
                    ('Building Levels', f"{building[5]}" if building[5] else 'N/A')
                ]
                
                for i, (metric, value) in enumerate(metrics):
                    table.cell(i+1, 0).text = metric
                    table.cell(i+1, 1).text = value
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            
            # Read file content
            with open(temp_file.name, 'rb') as f:
                docx_content = f.read()
            
            # Cleanup
            os.unlink(temp_file.name)
            
            return docx_content
            
        except Exception as e:
            st.error(f"Error generating Word report: {str(e)}")
            return None

def create_download_links(pdf_content, docx_content, project_id):
    """Create download links for generated reports"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if pdf_content:
        b64_pdf = base64.b64encode(pdf_content).decode()
        pdf_filename = f"BIPV_Comprehensive_Report_Project_{project_id}_{timestamp}.pdf"
        st.markdown(f"""
        <a href="data:application/pdf;base64,{b64_pdf}" download="{pdf_filename}">
            <button style="background-color:#2e75b6;color:white;padding:10px 20px;border:none;border-radius:5px;cursor:pointer;">
                📄 Download PDF Report
            </button>
        </a>
        """, unsafe_allow_html=True)
    
    if docx_content:
        b64_docx = base64.b64encode(docx_content).decode()
        docx_filename = f"BIPV_Comprehensive_Report_Project_{project_id}_{timestamp}.docx"
        st.markdown(f"""
        <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="{docx_filename}">
            <button style="background-color:#2e75b6;color:white;padding:10px 20px;border:none;border-radius:5px;cursor:pointer;">
                📝 Download Word Report
            </button>
        </a>
        """, unsafe_allow_html=True)